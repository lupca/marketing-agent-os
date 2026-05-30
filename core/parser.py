# core/parser.py
"""
Document Parser & Semantic Chunker.

Hỗ trợ: PDF, TXT, DOCX, XLSX
Chunking: RecursiveCharacterTextSplitter (Semantic Boundary)
  chunk_size=1000, chunk_overlap=200 — theo quyết định CTO v3
"""
import os
import logging
from typing import List

logger = logging.getLogger("core_parser")


# ============================================================
# TEXT EXTRACTION
# ============================================================
def extract_text_from_file(file_path: str) -> str:
    """
    Extract plain text từ các định dạng file được hỗ trợ.
    Supported: .txt, .pdf, .docx, .xlsx
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File không tồn tại: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        return _extract_txt(file_path)
    elif ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".xlsx":
        return _extract_xlsx(file_path)
    else:
        raise ValueError(f"Định dạng file không được hỗ trợ: '{ext}'. Chấp nhận: .txt, .pdf, .docx, .xlsx")


def _extract_txt(file_path: str) -> str:
    logger.info(f"[parser] Extracting TXT: {file_path}")
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _extract_pdf(file_path: str) -> str:
    logger.info(f"[parser] Extracting PDF: {file_path}")
    from pypdf import PdfReader
    try:
        reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"[Trang {i+1}]\n{text}")
        return "\n\n".join(pages)
    except Exception as e:
        logger.error(f"[parser] Lỗi đọc PDF: {e}")
        raise


def _extract_docx(file_path: str) -> str:
    """Extract từ DOCX, giữ nguyên cấu trúc heading và paragraph."""
    logger.info(f"[parser] Extracting DOCX: {file_path}")
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Giữ heading để chunk chia đúng ranh giới ngữ nghĩa
                if para.style.name.startswith("Heading"):
                    paragraphs.append(f"\n## {text}\n")
                else:
                    paragraphs.append(text)

        # Xử lý bảng trong DOCX
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n" + "\n".join(rows) + "\n")

        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"[parser] Lỗi đọc DOCX: {e}")
        raise


def _extract_xlsx(file_path: str) -> str:
    """Extract từ XLSX — chuyển mỗi sheet thành bảng text."""
    logger.info(f"[parser] Extracting XLSX: {file_path}")
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheets_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = [f"## Sheet: {sheet_name}"]
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                # Bỏ qua hàng hoàn toàn rỗng
                if any(c.strip() for c in cells):
                    rows_text.append(" | ".join(cells))
            if len(rows_text) > 1:  # Có data ngoài header
                sheets_text.append("\n".join(rows_text))

        wb.close()
        return "\n\n".join(sheets_text)
    except Exception as e:
        logger.error(f"[parser] Lỗi đọc XLSX: {e}")
        raise


# ============================================================
# SEMANTIC CHUNKING (Quyết định CTO v3)
# ============================================================
def semantic_chunk_text(
    text: str,
    chunk_size: int = None,
    chunk_overlap: int = None,
    separators: list = None,
) -> List[str]:
    """
    Chia text thành chunks theo ranh giới ngữ nghĩa (Semantic Boundary).

    Sử dụng RecursiveCharacterTextSplitter từ LangChain:
    - Thử chia theo paragraph (\\n\\n) trước
    - Nếu chunk vẫn quá lớn → chia theo dòng (\\n)
    - Nếu vẫn quá lớn → chia theo câu (., ?, !)
    - Cuối cùng → chia theo từ (space)

    Tham số mặc định theo config.settings (CTO v3):
      chunk_size=1000, chunk_overlap=200

    Returns:
        List[str]: Danh sách nội dung các chunk (không có metadata)
    """
    if not text or not text.strip():
        return []

    # Import settings tại đây để tránh circular import
    from config.settings import (
        RAG_CHUNK_SIZE,
        RAG_CHUNK_OVERLAP,
        RAG_CHUNK_SEPARATORS,
    )

    _chunk_size    = chunk_size    or RAG_CHUNK_SIZE
    _chunk_overlap = chunk_overlap or RAG_CHUNK_OVERLAP
    _separators    = separators    or RAG_CHUNK_SEPARATORS

    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=_chunk_size,
            chunk_overlap=_chunk_overlap,
            separators=_separators,
            length_function=len,
            is_separator_regex=False,
        )

        chunks = splitter.split_text(text.strip())
        # Lọc bỏ chunks quá ngắn (< 10 ký tự — thường là page markers, whitespace)
        chunks = [c.strip() for c in chunks if c and len(c.strip()) >= 10]

        logger.info(
            f"[parser] semantic_chunk_text: {len(text)} chars → {len(chunks)} chunks "
            f"(size={_chunk_size}, overlap={_chunk_overlap})"
        )
        return chunks

    except ImportError:
        # Fallback nếu langchain-text-splitters chưa install (development)
        logger.warning("[parser] langchain_text_splitters not found — using char fallback")
        return _char_fallback_chunk(text, _chunk_size, _chunk_overlap)


def _char_fallback_chunk(text: str, chunk_size: int, overlap: int) -> List[str]:
    """Fallback chunker dựa trên ký tự — chỉ dùng khi langchain chưa cài."""
    chunks = []
    text = text.strip()
    text_len = len(text)
    start = 0
    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk = text[start:end].strip()
        if len(chunk) >= 10:
            chunks.append(chunk)
        if chunk_size <= overlap:
            break
        start += (chunk_size - overlap)
    logger.info(f"[parser] fallback_chunk: {text_len} chars → {len(chunks)} chunks")
    return chunks
